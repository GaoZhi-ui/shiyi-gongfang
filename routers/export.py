"""
一键导出路由

POST /api/v1/export/docx       — 导出为 docx（传章节 ID 列表或 "all"）
POST /api/v1/export/txt        — 导出为 txt
POST /api/v1/export/markdown   — 导出为 markdown 压缩包

导出文件存 export/ 临时目录，返回下载链接。
路径前缀 /api/v1/export
"""

import json, re, zipfile, uuid
from io import BytesIO
from pathlib import Path
from datetime import datetime, timezone
from typing import Literal
from fastapi import APIRouter, HTTPException
from fastapi.responses import FileResponse
from pydantic import BaseModel, Field

router = APIRouter(prefix="/export", tags=["export"])
BASE = Path(__file__).resolve().parent.parent
EXPORT_DIR = BASE / "export"

EXPORT_DIR.mkdir(parents=True, exist_ok=True)

MAX_EXPORT_FILES = 100


# ─── 共享：章节目录解析 ───

def _resolve_chapters_dir() -> Path:
    """从 config.yaml 读取章节目录路径，失败则回退到本地 chapters/"""
    yaml_path = BASE / "config.yaml"
    if yaml_path.exists():
        try:
            import yaml
            cfg = yaml.safe_load(yaml_path.read_text(encoding="utf-8"))
            active = cfg.get("active_project", "")
            proj = cfg.get("projects", {}).get(active, {})
            root = proj.get("root")
            chapters_rel = proj.get("chapters_dir", "chapters")
            if root:
                p = Path(root).expanduser().resolve() / chapters_rel
                if p.is_dir():
                    return p
        except Exception:
            pass
    local_dir = (BASE / "chapters").resolve()
    local_dir.mkdir(exist_ok=True)
    return local_dir


CHAPTERS_DIR = _resolve_chapters_dir()
ALLOWED_EXTENSIONS = {".md"}


# ─── 请求体模型 ───


class ExportBody(BaseModel):
    chapters: list[str] | Literal["all"] = Field(
        "all",
        description="要导出的章节列表，或 'all' 表示全部",
    )
    title: str | None = Field(
        None,
        description="导出文件标题（可选，默认使用项目名或当前日期）",
    )


# ─── 工具函数 ───


def _collect_chapters(
    chapters_param: list[str] | str,
) -> list[tuple[str, str]]:
    """
    根据参数收集章节文件内容。
    返回 [(filename, content), ...] 列表。
    """
    all_files: list[Path] = sorted(
        CHAPTERS_DIR.glob("*.md"),
        key=lambda p: p.stat().st_mtime,
    )

    if not all_files:
        raise HTTPException(400, detail={
            "code": "NO_CHAPTERS",
            "message": "当前项目没有章节文件，无法导出",
        })

    selected: list[Path] = []

    if chapters_param == "all":
        selected = all_files
    else:
        if not isinstance(chapters_param, list) or not chapters_param:
            raise HTTPException(400, detail={
                "code": "INVALID_PARAMETER",
                "message": "chapters 必须为 'all' 或非空文件名列表",
            })
        requested = set(chapters_param)
        for f in all_files:
            if f.name in requested:
                selected.append(f)

        not_found = requested - {f.name for f in all_files}
        if not_found:
            raise HTTPException(404, detail={
                "code": "CHAPTERS_NOT_FOUND",
                "message": f"以下章节不存在: {', '.join(sorted(not_found))}",
                "hint": "请先通过 GET /api/v1/chapters 查看可用章节列表",
            })

    if len(selected) > MAX_EXPORT_FILES:
        raise HTTPException(413, detail={
            "code": "TOO_MANY_FILES",
            "message": f"单次导出最多 {MAX_EXPORT_FILES} 个文件，当前选择了 {len(selected)} 个",
        })

    result: list[tuple[str, str]] = []
    for f in selected:
        result.append((f.name, f.read_text(encoding="utf-8", errors="replace")))

    return result


def _export_title(chapters_param: list[str] | str, custom_title: str | None) -> str:
    """生成导出文件标题"""
    if custom_title:
        return custom_title
    now = datetime.now(timezone.utc).astimezone()
    return f"写作助手_导出_{now.strftime('%Y%m%d_%H%M')}"


def _safe_filename(title: str, ext: str) -> str:
    """生成安全文件名"""
    safe = re.sub(r'[<>:"/\\|?*]', '_', title)
    safe = safe.strip().replace(' ', '_')
    if not safe:
        safe = "export"
    return f"{safe}{ext}"


# ─── 内部导出函数 ───


def _export_docx(chapters: list[tuple[str, str]], title: str) -> Path:
    """生成 docx 文件"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(500, detail={
            "code": "DEPENDENCY_MISSING",
            "message": "python-docx 未安装，请执行 pip install python-docx",
        })

    doc = Document()

    # 文档标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元数据
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        f"共 {len(chapters)} 章 | 导出时间: "
        f"{datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    for i, (filename, content) in enumerate(chapters):
        # 章节标题
        doc.add_heading(filename.replace(".md", ""), level=1)

        # 解析正文与日记
        parts = content.split("---", 1)
        body = parts[0].strip()
        diary = parts[1].strip() if len(parts) > 1 else ""

        # 正文
        for line in body.split("\n"):
            line = line.strip()
            if not line:
                continue
            if line.startswith("#"):
                level = min(len(line.split()[0]), 4)
                doc.add_heading(line.lstrip("#").strip(), level=level)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\. ", line):
                doc.add_paragraph(line, style="List Number")
            else:
                p = doc.add_paragraph(line)

        # 日记部分
        if diary:
            doc.add_paragraph("")  # spacing
            diary_label = doc.add_paragraph()
            run_lbl = diary_label.add_run("—— 章末日记 ——")
            run_lbl.bold = True
            run_lbl.font.size = Pt(10)
            run_lbl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            for line in diary.split("\n"):
                line = line.strip()
                if line:
                    p = doc.add_paragraph(line)
                    p.paragraph_format.left_indent = Cm(1)

        # 章节间分页（最后一章不分页）
        if i < len(chapters) - 1:
            doc.add_page_break()

    filename = _safe_filename(title, ".docx")
    filepath = EXPORT_DIR / filename
    doc.save(str(filepath))
    return filepath


def _export_txt(chapters: list[tuple[str, str]], title: str) -> Path:
    """生成 txt 文件"""
    lines: list[str] = []
    lines.append(f"{'='*60}")
    lines.append(f"{title:^60}")
    lines.append(f"{'='*60}")
    lines.append(f"共 {len(chapters)} 章")
    lines.append(f"导出时间: {datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M')}")
    lines.append("")
    lines.append(f"{'='*60}")
    lines.append("")

    for filename, content in chapters:
        lines.append("")
        lines.append(f"{'─'*40}")
        lines.append(f"  {filename.replace('.md', '')}")
        lines.append(f"{'─'*40}")
        lines.append("")

        parts = content.split("---", 1)
        lines.append(parts[0].strip())

        if len(parts) > 1 and parts[1].strip():
            lines.append("")
            lines.append("[章末日记]")
            lines.append(parts[1].strip())

        lines.append("")

    content = "\n".join(lines)
    filename = _safe_filename(title, ".txt")
    filepath = EXPORT_DIR / filename
    filepath.write_text(content, encoding="utf-8")
    return filepath


def _export_markdown_zip(chapters: list[tuple[str, str]], title: str) -> Path:
    """生成 markdown 压缩包"""
    buf = BytesIO()

    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # 添加一个元信息文件
        meta = (
            f"# {title}\n\n"
            f"导出时间: {datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M')}\n"
            f"章节数: {len(chapters)}\n\n"
            f"| # | 文件名 | 字数 |\n"
            f"|---|--------|------|\n"
        )
        for i, (filename, content) in enumerate(chapters, 1):
            cjk = len(re.findall(r'[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]', content))
            meta += f"| {i} | {filename} | {cjk} |\n"

        zf.writestr("README.md", meta.encode("utf-8"))

        # 添加各章节文件
        for filename, content in chapters:
            zf.writestr(filename, content.encode("utf-8"))

    filename = _safe_filename(title, ".zip")
    filepath = EXPORT_DIR / filename
    filepath.write_bytes(buf.getvalue())
    return filepath


# ─── 路由 ───


@router.post("/docx", status_code=201)
def export_docx(body: ExportBody):
    """导出为 DOCX 文档"""
    chapters = _collect_chapters(body.chapters)
    title = _export_title(body.chapters, body.title)
    filepath = _export_docx(chapters, title)
    return {
        "status": "ok",
        "format": "docx",
        "filename": filepath.name,
        "chapter_count": len(chapters),
        "download_url": f"/export/{filepath.name}",
    }


@router.post("/txt", status_code=201)
def export_txt(body: ExportBody):
    """导出为 TXT 纯文本"""
    chapters = _collect_chapters(body.chapters)
    title = _export_title(body.chapters, body.title)
    filepath = _export_txt(chapters, title)
    return {
        "status": "ok",
        "format": "txt",
        "filename": filepath.name,
        "chapter_count": len(chapters),
        "download_url": f"/export/{filepath.name}",
    }


@router.post("/markdown", status_code=201)
def export_markdown(body: ExportBody):
    """导出为 Markdown 压缩包"""
    chapters = _collect_chapters(body.chapters)
    title = _export_title(body.chapters, body.title)
    filepath = _export_markdown_zip(chapters, title)
    return {
        "status": "ok",
        "format": "zip",
        "filename": filepath.name,
        "chapter_count": len(chapters),
        "download_url": f"/export/{filepath.name}",
    }
