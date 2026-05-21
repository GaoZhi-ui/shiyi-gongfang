"""
Token化导出消费者

将 FormattedDocument（Token流中间表示）渲染为具体输出格式。
每个 Consumer 继承 BaseConsumer 并实现 consume() 方法。
"""

from __future__ import annotations
import re
import uuid
import zipfile
from abc import ABC, abstractmethod
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any

from .export_engine import FormattedDocument, ChapterToken


def _safe_filename(title: str, ext: str) -> str:
    """生成安全文件名"""
    safe = re.sub(r'[<>:"/\\|?*]', "_", title)
    safe = safe.strip().replace(" ", "_")
    if not safe:
        safe = "export"
    return f"{safe}{ext}"


# ─── 抽象基类 ───


class BaseConsumer(ABC):
    """导出消费者抽象基类。

    子类必须实现 consume(document) -> Any。
    """

    def __init__(self, export_dir: str | Path):
        self.export_dir = Path(export_dir)
        self.export_dir.mkdir(parents=True, exist_ok=True)

    @abstractmethod
    def consume(self, document: FormattedDocument) -> Any:
        """消费 FormattedDocument，生成目标格式文件。"""
        ...


# ─── 辅助：从 Token 流重新渲染为 Markdown ───

def _tokens_to_markdown(tokens: list[ChapterToken], include_diary_label: bool = True) -> str:
    """将Token列表渲染为Markdown文本"""
    lines: list[str] = []
    diary_mode = False
    for token in tokens:
        if token.type == "heading":
            prefix = "#" * token.level
            lines.append(f"{prefix} {token.content[0]}")
            lines.append("")
        elif token.type == "paragraph":
            lines.extend(token.content)
            lines.append("")
        elif token.type == "list":
            for item in token.content:
                lines.append(f"- {item}")
            lines.append("")
        elif token.type == "blank":
            lines.append("")
        elif token.type == "diary":
            if include_diary_label and not diary_mode:
                # Insert the diary separator
                if lines and lines[-1] != "":
                    lines.append("")
                lines.append("---")
                lines.append("")
                diary_mode = True
            lines.extend(token.content)
            lines.append("")
    # 去掉末尾多余的换行
    while lines and lines[-1] == "":
        lines.pop()
    lines.append("")
    return "\n".join(lines)


# ─── Docx 消费者 ───


class DocxConsumer(BaseConsumer):
    """将 FormattedDocument 渲染为 DOCX 文件。"""

    def consume(self, document: FormattedDocument) -> Path:
        """生成 .docx 文件，返回文件路径"""
        try:
            from docx import Document
            from docx.shared import Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            raise RuntimeError(
                "python-docx 未安装，请执行 pip install python-docx"
            )

        doc = Document()

        # 文档标题
        title_para = doc.add_heading(document.title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 元数据
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        now_str = datetime.now(timezone.utc).astimezone().strftime("%Y-%m-%d %H:%M")
        run = meta.add_run(
            f"共 {document.chapter_count} 章 | 导出时间: {now_str}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        doc.add_page_break()

        for i, (filename, tokens) in enumerate(document.chapters):
            # 章节标题
            chapter_title = filename.replace(".md", "")
            doc.add_heading(chapter_title, level=1)

            # 遍历Token渲染
            diary_mode = False
            for token in tokens:
                if token.type == "heading":
                    level = min(token.level, 4)
                    doc.add_heading(token.content[0], level=level)

                elif token.type == "paragraph":
                    for line in token.content:
                        stripped = line.strip()
                        if stripped:
                            doc.add_paragraph(stripped)

                elif token.type == "list":
                    for item in token.content:
                        doc.add_paragraph(item, style="List Bullet")

                elif token.type == "blank":
                    pass  # 跳过，DOCX段落间距自动处理

                elif token.type == "diary":
                    if not diary_mode:
                        diary_mode = True
                        doc.add_paragraph("")  # spacing
                        diary_label = doc.add_paragraph()
                        run_lbl = diary_label.add_run("—— 章末日记 ——")
                        run_lbl.bold = True
                        run_lbl.font.size = Pt(10)
                        run_lbl.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

                    for line in token.content:
                        stripped = line.strip()
                        if stripped:
                            p = doc.add_paragraph(stripped)
                            p.paragraph_format.left_indent = Cm(1)

            # 章节间分页
            if i < len(document.chapters) - 1:
                doc.add_page_break()

        filename = _safe_filename(document.title, ".docx")
        filepath = self.export_dir / filename
        doc.save(str(filepath))
        return filepath


# ─── Txt 消费者 ───


class TxtConsumer(BaseConsumer):
    """将 FormattedDocument 渲染为 TXT 纯文本文件。"""

    def __init__(self, export_dir: str | Path, line_width: int = 60):
        super().__init__(export_dir)
        self.line_width = line_width

    def consume(self, document: FormattedDocument) -> Path:
        """生成 .txt 文件，返回文件路径"""
        now_str = datetime.now(timezone.utc).astimezone().strftime("%Y-%m-%d %H:%M")
        w = self.line_width
        lines: list[str] = []

        # 文档标题横幅
        lines.append(f"{'=' * w}")
        lines.append(f"{document.title:^{w}}")
        lines.append(f"{'=' * w}")
        lines.append(f"共 {document.chapter_count} 章")
        lines.append(f"导出时间: {now_str}")
        lines.append("")
        lines.append(f"{'=' * w}")
        lines.append("")

        for filename, tokens in document.chapters:
            chapter_title = filename.replace(".md", "")
            lines.append("")
            lines.append(f"{'─' * 40}")
            lines.append(f"  {chapter_title}")
            lines.append(f"{'─' * 40}")
            lines.append("")

            diary_mode = False
            for token in tokens:
                if token.type == "heading":
                    prefix = "#" * token.level
                    for item in token.content:
                        lines.append(f"{prefix} {item}")
                    lines.append("")

                elif token.type == "paragraph":
                    for line_text in token.content:
                        stripped = line_text.strip()
                        if stripped:
                            lines.append(stripped)
                    lines.append("")

                elif token.type == "list":
                    for item in token.content:
                        lines.append(f"  - {item}")
                    lines.append("")

                elif token.type == "blank":
                    lines.append("")

                elif token.type == "diary":
                    if not diary_mode:
                        diary_mode = True
                        lines.append("")
                        lines.append("[章末日记]")
                    for line_text in token.content:
                        stripped = line_text.strip()
                        if stripped:
                            lines.append(f"  {stripped}")
                    lines.append("")

        content = "\n".join(lines)
        filename = _safe_filename(document.title, ".txt")
        filepath = self.export_dir / filename
        filepath.write_text(content, encoding="utf-8")
        return filepath


# ─── PDF 消费者 ───


def _find_cjk_font() -> str | None:
    """查找系统中可用的 CJK 字体，按平台依次尝试。"""
    import platform as _platform
    from pathlib import Path

    system = _platform.system()

    candidates: list[str] = []
    if system == "Windows":
        candidates = [
            r"C:\Windows\Fonts\msyh.ttc",        # 微软雅黑
            r"C:\Windows\Fonts\simsun.ttc",       # 宋体
            r"C:\Windows\Fonts\simhei.ttf",       # 黑体
            r"C:\Windows\Fonts\deng.ttf",         # 等线
            r"C:\Windows\Fonts\yahei.ttf",        # 雅黑（备选）
        ]
    elif system == "Darwin":
        candidates = [
            "/System/Library/Fonts/PingFang.ttc",
            "/System/Library/Fonts/STHeiti Light.ttc",
            "/Library/Fonts/Arial Unicode.ttf",
        ]
    else:  # Linux
        candidates = [
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/noto-cjk/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
            "/usr/share/fonts/wqy-zenhei/wqy-zenhei.ttc",
        ]

    for path in candidates:
        p = Path(path)
        if p.exists():
            return str(p)
    return None


class PdfConsumer(BaseConsumer):
    """将 FormattedDocument 渲染为 PDF 文件。

    使用 fpdf2 库（纯 Python，无系统依赖）。
    包含标题页、章节标题、正文、自动分页。
    """

    def __init__(
        self,
        export_dir: str | Path,
        font_path: str | None = None,
        font_name: str = "CJK",
        title_font_size: int = 24,
        chapter_font_size: int = 16,
        body_font_size: int = 11,
    ):
        super().__init__(export_dir)
        self.font_name = font_name
        self.title_font_size = title_font_size
        self.chapter_font_size = chapter_font_size
        self.body_font_size = body_font_size

        # 解析字体路径
        self.font_path = font_path or _find_cjk_font()
        if not self.font_path:
            raise RuntimeError(
                "未找到系统 CJK 字体。请通过 font_path 参数指定中文字体路径，"
                "或安装 Noto Sans CJK / 微软雅黑等字体。"
            )

    def _init_pdf(self) -> any:
        """初始化 FPDF 实例并注册字体"""
        from fpdf import FPDF

        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=20)

        # TTC 字体集需要指定索引
        font_path_lower = self.font_path.lower()
        if font_path_lower.endswith(".ttc"):
            pdf.add_font(self.font_name, "", self.font_path, collection_font_number=0)
        else:
            pdf.add_font(self.font_name, "", self.font_path)
        return pdf

    def _add_title_page(self, pdf: any, document: any) -> None:
        """生成标题页"""
        from datetime import datetime, timezone

        pdf.add_page()
        pdf.ln(50)  # 垂直居中

        # 书名
        pdf.set_font(self.font_name, size=self.title_font_size)
        pdf.cell(w=0, text=document.title, align="C", new_x="LMARGIN", new_y="NEXT")

        pdf.ln(10)

        # 元信息
        pdf.set_font(self.font_name, size=11)
        now_str = datetime.now(timezone.utc).astimezone().strftime("%Y-%m-%d %H:%M")
        pdf.cell(w=0, text=f"共 {document.chapter_count} 章", align="C", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(w=0, text=f"导出时间: {now_str}", align="C", new_x="LMARGIN", new_y="NEXT")

    def _render_tokens(self, pdf: any, tokens: list[ChapterToken]) -> None:
        """将 Token 列表渲染到 PDF"""
        diary_mode = False

        for token in tokens:
            if token.type == "heading":
                if token.level <= 2:
                    pdf.set_font(self.font_name, size=self.chapter_font_size)
                    pdf.ln(4)
                    for item in token.content:
                        pdf.multi_cell(w=0, text=item)
                    pdf.ln(3)
                else:
                    pdf.set_font(self.font_name, size=self.body_font_size + 2)
                    pdf.ln(2)
                    for item in token.content:
                        pdf.multi_cell(w=0, text=item)
                    pdf.ln(2)

            elif token.type == "paragraph":
                pdf.set_font(self.font_name, size=self.body_font_size)
                for line in token.content:
                    stripped = line.strip()
                    if stripped:
                        pdf.multi_cell(w=0, text=stripped)
                        pdf.ln(1)
                pdf.ln(2)

            elif token.type == "list":
                pdf.set_font(self.font_name, size=self.body_font_size)
                for item in token.content:
                    pdf.multi_cell(w=0, text=f"  · {item}")
                    pdf.ln(1)
                pdf.ln(2)

            elif token.type == "blank":
                pass

            elif token.type == "diary":
                if not diary_mode:
                    diary_mode = True
                    pdf.ln(3)
                    pdf.set_font(self.font_name, size=self.body_font_size)
                    pdf.set_text_color(100, 100, 100)
                    pdf.multi_cell(w=0, text="—— 章末日记 ——")
                    pdf.ln(2)

                for line in token.content:
                    stripped = line.strip()
                    if stripped:
                        pdf.set_font(self.font_name, size=self.body_font_size)
                        pdf.set_text_color(80, 80, 80)
                        pdf.multi_cell(w=0, text=stripped)
                        pdf.ln(1)

                pdf.set_text_color(0, 0, 0)

    def consume(self, document: FormattedDocument) -> Path:
        """生成 .pdf 文件，返回文件路径"""
        pdf = self._init_pdf()

        # 标题页
        self._add_title_page(pdf, document)

        # 各章节
        for i, (filename, tokens) in enumerate(document.chapters):
            pdf.add_page()

            # 章节标题
            chapter_title = filename.replace(".md", "")
            pdf.set_font(self.font_name, size=self.chapter_font_size)
            pdf.multi_cell(w=0, text=chapter_title)
            pdf.ln(4)

            # 渲染内容
            self._render_tokens(pdf, tokens)

        filename = _safe_filename(document.title, ".pdf")
        filepath = self.export_dir / filename
        pdf.output(str(filepath))
        return filepath


# ─── EPUB 消费者 ───


class EpubConsumer(BaseConsumer):
    """将 FormattedDocument 渲染为 EPUB 电子书。

    使用 ebooklib 库生成标准 EPUB 文件，包含:
      - 标题页
      - 章节标题 + 正文 + 章末日记
      - NCX 自动目录
      - 样式表
    """

    def __init__(
        self,
        export_dir: str | Path,
        author: str = "写作助手工坊",
    ):
        super().__init__(export_dir)
        self.author = author

    def _tokens_to_html(self, tokens: list[ChapterToken]) -> str:
        """将 Token 流渲染为 HTML 片段"""
        html_parts: list[str] = []
        diary_mode = False

        for token in tokens:
            if token.type == "heading":
                tag = f"h{min(token.level, 4)}"
                for item in token.content:
                    html_parts.append(f"<{tag}>{_escape_html(item)}</{tag}>")

            elif token.type == "paragraph":
                for line in token.content:
                    stripped = line.strip()
                    if stripped:
                        html_parts.append(f"<p>{_escape_html(stripped)}</p>")

            elif token.type == "list":
                html_parts.append("<ul>")
                for item in token.content:
                    html_parts.append(f"<li>{_escape_html(item)}</li>")
                html_parts.append("</ul>")

            elif token.type == "blank":
                pass

            elif token.type == "diary":
                if not diary_mode:
                    diary_mode = True
                    html_parts.append('<hr class="diary-sep" />')
                    html_parts.append(
                        '<p class="diary-label">—— 章末日记 ——</p>'
                    )
                for line in token.content:
                    stripped = line.strip()
                    if stripped:
                        html_parts.append(
                            f'<p class="diary">{_escape_html(stripped)}</p>'
                        )

        return "\n".join(html_parts)

    def _make_title_page(self, document: FormattedDocument) -> str:
        """生成 EPUB 标题页 HTML"""
        now_str = datetime.now(timezone.utc).astimezone().strftime(
            "%Y-%m-%d %H:%M"
        )
        return f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>{_escape_html(document.title)}</title></head>
<body>
  <div style="text-align:center;margin-top:30vh;">
    <h1 style="font-size:2em;">{_escape_html(document.title)}</h1>
    <p style="color:#666;">作者: {_escape_html(self.author)}</p>
    <p style="color:#999;font-size:0.9em;">共 {document.chapter_count} 章</p>
    <p style="color:#999;font-size:0.9em;">导出时间: {now_str}</p>
  </div>
</body>
</html>"""

    def _make_chapter_html(
        self,
        chapter_title: str,
        tokens: list[ChapterToken],
        uid: str,
    ) -> str:
        """生成单章 HTML"""
        body_html = self._tokens_to_html(tokens)
        return f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head>
  <title>{_escape_html(chapter_title)}</title>
  <link rel="stylesheet" type="text/css" href="style.css"/>
</head>
<body>
  <h1 class="chapter-title">{_escape_html(chapter_title)}</h1>
  {body_html}
</body>
</html>"""

    def _make_css(self) -> str:
        """生成 EPUB 样式表"""
        return """
@page { margin: 1em 1.5em; }
body { font-family: "Noto Serif", "Source Han Serif", serif;
       line-height: 1.8; color: #333; }
h1.chapter-title { text-align: center; font-size: 1.6em;
                   margin: 2em 0 1.5em; }
h2 { font-size: 1.3em; margin: 1.5em 0 0.8em; }
h3, h4 { font-size: 1.1em; margin: 1.2em 0 0.6em; }
p { text-indent: 2em; margin: 0.5em 0; }
ul { margin: 0.5em 0 0.5em 2em; }
li { margin: 0.3em 0; }
hr.diary-sep { border: none; border-top: 1px dashed #ccc;
               margin: 2em 0 1em; }
p.diary-label { text-align: center; color: #888;
                font-size: 0.9em; }
p.diary { color: #666; font-style: italic;
           text-indent: 2em; }
"""

    def consume(self, document: FormattedDocument) -> Path:
        """生成 .epub 文件，返回文件路径"""
        try:
            import ebooklib
            from ebooklib import epub
        except ImportError:
            raise RuntimeError(
                "ebooklib 未安装，请执行 pip install ebooklib"
            )

        book = epub.EpubBook()

        # 元数据
        book.set_identifier(
            f"urn:uuid:{uuid.uuid4().hex[:32]}"
        )
        book.set_title(document.title)
        book.set_language("zh-CN")
        book.add_author(self.author)

        # 样式表
        css_content = self._make_css()
        css_item = epub.EpubItem(
            uid="style",
            file_name="style.css",
            media_type="text/css",
            content=css_content.encode("utf-8"),
        )
        book.add_item(css_item)

        # 标题页
        title_html = self._make_title_page(document)
        title_item = epub.EpubHtml(
            uid="title",
            file_name="title.xhtml",
            content=title_html.encode("utf-8"),
        )
        title_item.add_item(css_item)
        book.add_item(title_item)

        # 章节
        spine = ["nav", title_item]
        toc = []

        for i, (filename, tokens) in enumerate(document.chapters):
            chapter_title = filename.replace(".md", "")
            uid = f"chap_{i+1:03d}"

            chap_html = self._make_chapter_html(
                chapter_title, tokens, uid
            )
            chap_item = epub.EpubHtml(
                uid=uid,
                file_name=f"chapter_{i+1:03d}.xhtml",
                content=chap_html.encode("utf-8"),
            )
            chap_item.add_item(css_item)
            book.add_item(chap_item)

            spine.append(chap_item)
            toc.append(epub.Link(
                f"chapter_{i+1:03d}.xhtml",
                chapter_title,
                uid,
            ))

        # NCX 目录 + 导航
        book.toc = toc
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        book.spine = spine

        # 写入文件
        filename = _safe_filename(document.title, ".epub")
        filepath = self.export_dir / filename
        epub.write_epub(str(filepath), book)
        return filepath


# ─── 电子书编译消费者（多章合并 + 封面 + 目录） ───


class EbookConsumer(BaseConsumer):
    """合并多章为完整电子书，含封面页和完整目录。

    输出为 EPUB 格式，复用 EpubConsumer 的渲染逻辑。
    额外功能：
      - 自动生成封面页（标题 + 作者 + 日期）
      - 完整章节目录
      - 多章合并为一个文件
    """

    def __init__(
        self,
        export_dir: str | Path,
        author: str = "写作助手工坊",
        cover_subtitle: str | None = None,
    ):
        super().__init__(export_dir)
        self.author = author
        self.cover_subtitle = cover_subtitle

    def consume(self, document: FormattedDocument) -> Path:
        """生成含封面和目录的 .epub 文件，返回文件路径"""
        # 委派给 EpubConsumer，但先注入封面页
        # 用 EpubConsumer 的渲染管线
        epub_consumer = EpubConsumer(
            export_dir=self.export_dir,
            author=self.author,
        )

        # 用 FormattedDocument 的 all_tokens() 配合自定义封面
        return self._build_ebook(document, epub_consumer)

    def _build_ebook(
        self,
        document: FormattedDocument,
        epub_consumer: EpubConsumer,
    ) -> Path:
        """构建完整的电子书"""
        try:
            import ebooklib
            from ebooklib import epub
        except ImportError:
            raise RuntimeError(
                "ebooklib 未安装，请执行 pip install ebooklib"
            )

        book = epub.EpubBook()

        # 元数据
        book.set_identifier(
            f"urn:uuid:{uuid.uuid4().hex[:32]}"
        )
        book.set_title(document.title)
        book.set_language("zh-CN")
        book.add_author(self.author)

        # 样式表
        css_content = epub_consumer._make_css()
        css_item = epub.EpubItem(
            uid="style",
            file_name="style.css",
            media_type="text/css",
            content=css_content.encode("utf-8"),
        )
        book.add_item(css_item)

        # ── 封面页 ──
        now_str = datetime.now(timezone.utc).astimezone().strftime(
            "%Y-%m-%d"
        )
        subtitle_html = ""
        if self.cover_subtitle:
            subtitle_html = (
                f'<p style="color:#888;margin-top:0.5em;'
                f'font-size:1.1em;">{_escape_html(self.cover_subtitle)}</p>'
            )
        cover_html = f"""<?xml version="1.0" encoding="utf-8"?>
<!DOCTYPE html>
<html xmlns="http://www.w3.org/1999/xhtml">
<head><title>{_escape_html(document.title)}</title></head>
<body>
  <div style="text-align:center;margin-top:25vh;">
    <h1 style="font-size:2.2em;letter-spacing:0.1em;">{_escape_html(document.title)}</h1>
    {subtitle_html}
    <p style="color:#666;margin-top:2em;">作者: {_escape_html(self.author)}</p>
    <p style="color:#999;font-size:0.85em;">{now_str}</p>
    <hr style="width:30%;margin:2em auto;border:none;border-top:1px solid #ddd;" />
    <p style="color:#aaa;font-size:0.85em;">共 {document.chapter_count} 章</p>
  </div>
</body>
</html>"""

        cover_item = epub.EpubHtml(
            uid="cover",
            file_name="cover.xhtml",
            content=cover_html.encode("utf-8"),
        )
        cover_item.add_item(css_item)
        book.add_item(cover_item)

        # ── 目录页 ──
        toc_lines = [
            '<?xml version="1.0" encoding="utf-8"?>',
            '<!DOCTYPE html>',
            '<html xmlns="http://www.w3.org/1999/xhtml">',
            '<head><title>目录</title>',
            '<link rel="stylesheet" type="text/css" href="style.css"/>',
            '</head><body>',
            '<h1 style="text-align:center;margin:2em 0 1.5em;">目录</h1>',
            '<ul style="list-style:none;padding:0;">',
        ]
        for i, (filename, _) in enumerate(document.chapters):
            chap_title = filename.replace(".md", "")
            html_fn = f"chapter_{i+1:03d}.xhtml"
            toc_lines.append(
                f'<li style="margin:0.6em 0;">'
                f'<a href="{html_fn}" '
                f'style="text-decoration:none;color:#333;'
                f'font-size:1.1em;">{_escape_html(chap_title)}</a></li>'
            )
        toc_lines.append("</ul></body></html>")

        toc_item = epub.EpubHtml(
            uid="toc_page",
            file_name="toc.xhtml",
            content="\n".join(toc_lines).encode("utf-8"),
        )
        toc_item.add_item(css_item)
        book.add_item(toc_item)

        # ── 章节内容 ──
        spine = ["nav", cover_item, toc_item]
        epub_toc: list = []

        for i, (filename, tokens) in enumerate(document.chapters):
            chap_title = filename.replace(".md", "")
            uid = f"chap_{i+1:03d}"

            chap_html = epub_consumer._make_chapter_html(
                chap_title, tokens, uid
            )
            chap_item = epub.EpubHtml(
                uid=uid,
                file_name=f"chapter_{i+1:03d}.xhtml",
                content=chap_html.encode("utf-8"),
            )
            chap_item.add_item(css_item)
            book.add_item(chap_item)

            spine.append(chap_item)
            epub_toc.append(epub.Link(
                f"chapter_{i+1:03d}.xhtml",
                chap_title,
                uid,
            ))

        # NCX + Nav
        book.toc = epub_toc
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())

        book.spine = spine

        # 写入
        filename = _safe_filename(document.title, ".epub")
        filepath = self.export_dir / filename
        epub.write_epub(str(filepath), book)
        return filepath


# ─── 辅助：HTML 转义 ───


def _escape_html(text: str) -> str:
    """转义 HTML 特殊字符"""
    text = text.replace("&", "&amp;")
    text = text.replace("<", "&lt;")
    text = text.replace(">", "&gt;")
    text = text.replace('"', "&quot;")
    text = text.replace("'", "&#39;")
    return text


# ─── Markdown 消费者 ───


class MarkdownConsumer(BaseConsumer):
    """将 FormattedDocument 渲染为 Markdown 源码压缩包。

    以独立 .md 文件 + README.md 元信息的方式打包为 ZIP。
    """

    def consume(self, document: FormattedDocument) -> Path:
        """生成 .zip 压缩包，返回文件路径"""
        buf = BytesIO()

        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
            # README.md 元信息
            readme_lines = [
                f"# {document.title}",
                "",
                f"导出时间: {datetime.now(timezone.utc).astimezone().strftime('%Y-%m-%d %H:%M')}",
                f"章节数: {document.chapter_count}",
                "",
                "| # | 文件名 | 字数 |",
                "|---|--------|------|",
            ]
            for i, (filename, tokens) in enumerate(document.chapters, 1):
                word_count = sum(
                    len(token.text)
                    for token in tokens
                )
                readme_lines.append(f"| {i} | {filename} | {word_count} |")

            zf.writestr("README.md", "\n".join(readme_lines).encode("utf-8"))

            # 各章节 Markdown 文件
            for filename, tokens in document.chapters:
                md_content = _tokens_to_markdown(tokens, include_diary_label=True)
                zf.writestr(filename, md_content.encode("utf-8"))

        filename = _safe_filename(document.title, ".zip")
        filepath = self.export_dir / filename
        filepath.write_bytes(buf.getvalue())
        return filepath
